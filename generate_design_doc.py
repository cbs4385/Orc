#!/usr/bin/env python3
"""Generate the Replayability Improvements Design Document as a .docx file."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# -- Style setup --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x2E, 0x4A, 0x1E)  # Dark green, orcish

def add_table(headers, rows):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 3'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()  # spacer
    return table

# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('Defender of the Orcish Marches', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_heading('Replayability Improvements\nDesign Document', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Version 1.0 | February 2026\n').bold = True
meta.add_run('Status: Proposal / Design Phase')
doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (manual)
# ============================================================
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Context & Problem Statement',
    '2. System 1: Achievements / Milestones',
    '3. System 2: Unlockable Mutators',
    '4. System 3: Meta-Progression (Legacy Ranks)',
    '5. System 4: Statistics Dashboard',
    '6. System Interconnections',
    '7. Implementation Order',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')
doc.add_page_break()

# ============================================================
# 1. CONTEXT
# ============================================================
doc.add_heading('1. Context & Problem Statement', level=1)

doc.add_heading('Current State', level=2)
doc.add_paragraph(
    'Defender of the Orcish Marches is an endless tower-defense survival game. '
    'Players defend a fortress against escalating orc waves across day/night cycles, '
    'managing gold economy, menial workers, defenders, and ballista weaponry.'
)
doc.add_paragraph('The game currently features:')
current_features = [
    'A composite scoring system and run history leaderboard (top 20 runs)',
    '4 difficulty levels (Easy / Normal / Hard / Nightmare)',
    '12 random daily events providing per-day modifiers',
    '6 enemy types unlocking progressively (days 1-10)',
    '4 defender types and various upgrades',
]
for f in current_features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('The Problem', level=2)
doc.add_paragraph(
    'Once a player has experienced all enemy types (by day 10) and tried each difficulty, '
    'there is little motivation to replay. Score-chasing alone is not enough for most players. '
    'There are no persistent goals, no unlockables, no way to vary the gameplay experience '
    'between runs, and no lifetime stat tracking to show long-term improvement.'
)

doc.add_heading('Goal', level=2)
doc.add_paragraph(
    'Design four interconnected replayability systems that give players reasons to come back:'
)
goals = [
    ('Short-term goals', 'Achievements'),
    ('Gameplay variety', 'Mutators'),
    ('Long-term growth', 'Meta-Progression'),
    ('Self-reflection', 'Statistics'),
]
for label, system in goals:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{label}: ').bold = True
    p.add_run(system)

doc.add_page_break()

# ============================================================
# 2. ACHIEVEMENTS
# ============================================================
doc.add_heading('2. System 1: Achievements / Milestones', level=1)

doc.add_heading('Overview', level=2)
doc.add_paragraph(
    'A persistent achievement system with 20 achievements across 5 categories. '
    'Each has a bronze / silver / gold tier providing escalating challenge. '
    'Achievements persist across runs via PlayerPrefs JSON storage.'
)

# -- Survival --
doc.add_heading('Survival Achievements (4)', level=2)
add_table(
    ['Achievement', 'Bronze', 'Silver', 'Gold'],
    [
        ['Stand Your Ground\nSurvive days', '5 days', '15 days', '30 days'],
        ['Iron Walls\nFinish a day with all walls above 75% HP', '1 time', '5 times', '15 times'],
        ['Nightwatch\nSurvive nights without taking wall damage', '3 nights', '10 nights', '25 nights'],
        ['Against All Odds\nSurvive days on Hard or Nightmare', '3 days', '10 days', '20 days'],
    ]
)

# -- Combat --
doc.add_heading('Combat Achievements (4)', level=2)
add_table(
    ['Achievement', 'Bronze', 'Silver', 'Gold'],
    [
        ['Orc Slayer\nKill enemies (cumulative across runs)', '100', '500', '2,000'],
        ['Boss Hunter\nSlay War Bosses', '1', '5', '15'],
        ['Sharpshooter\nKill enemies with ballista in a single day', '10', '25', '50'],
        ['Exterminator\nKill every enemy type in a single run', '4 types', '5 types', 'All 6 types'],
    ]
)

# -- Economy --
doc.add_heading('Economy Achievements (4)', level=2)
add_table(
    ['Achievement', 'Bronze', 'Silver', 'Gold'],
    [
        ['Treasure Hoarder\nEarn gold in a single run', '200g', '500g', '1,000g'],
        ['Efficient Commander\nFinish a day with 0 idle menials', '3 times', '10 times', '20 times'],
        ['Golden Age\nEarn gold (cumulative across runs)', '1,000g', '5,000g', '20,000g'],
        ['Penny Pincher\nSurvive without spending gold', 'Day 5', 'Day 8', 'Day 12'],
    ]
)

# -- Defenders --
doc.add_heading('Defender Achievements (4)', level=2)
add_table(
    ['Achievement', 'Bronze', 'Silver', 'Gold'],
    [
        ['Army Builder\nHave living defenders simultaneously', '5', '10', '20'],
        ['Specialist\nSurvive hiring only one defender type', 'Any type', 'On Normal+', 'On Hard+'],
        ['Engineering Corps\nRepair walls (cumulative HP repaired)', '500 HP', '2,000 HP', '10,000 HP'],
        ['Combined Arms\nHave all 4 defender types alive simultaneously', '1 time', '5 times', '15 times'],
    ]
)

# -- Special --
doc.add_heading('Special Achievements (4)', level=2)
add_table(
    ['Achievement', 'Bronze', 'Silver', 'Gold'],
    [
        ['Refugee Savior\nRefugees successfully reach fortress', '10', '50', '200'],
        ['Deforester\nClear vegetation tiles', '20', '100', '500'],
        ['Nightmare Survivor\nSurvive days on Nightmare difficulty', '1 day', '5 days', '10 days'],
        ['Score Legend\nReach composite score', '5,000', '15,000', '50,000'],
    ]
)

doc.add_heading('Data Model', level=2)
doc.add_paragraph(
    'Each achievement stores: id (string), name, description, category, '
    'bronze/silver/gold threshold values, a flag for cumulative vs single-run tracking, '
    'and an optional mutator unlock ID for the gold tier.'
)

doc.add_heading('Persistence', level=2)
items = [
    'Store as JSON in PlayerPrefs key "Achievements"',
    'Dictionary mapping achievement ID to current progress value and highest tier reached',
    'Updated at end of each run (in GameOverScreen) and during gameplay for cumulative stats',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('UI Integration', level=2)
ui_items = [
    ('Game Over Screen:', 'Show 1-3 newly earned achievements below the stats, with tier badge color (bronze/silver/gold)'),
    ('Main Menu:', 'New "Achievements" button opens a scrollable panel showing all 20 achievements grouped by category, with progress bars and tier indicators'),
    ('In-Game Toast:', 'Small notification in bottom-right when an achievement tier is earned mid-run'),
]
for label, desc in ui_items:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(label).bold = True
    p.add_run(f' {desc}')

doc.add_heading('Files to Create / Modify', level=2)
files = [
    ('New:', 'AchievementManager.cs - Singleton, tracks progress, checks thresholds, fires events'),
    ('New:', 'AchievementUI.cs - Main menu achievement panel (programmatic UI)'),
    ('Modify:', 'GameOverScreen.cs - Add achievement display section'),
    ('Modify:', 'RunStatsTracker.cs - Feed stats to AchievementManager at run end'),
    ('Modify:', 'MainMenuManager.cs - Add Achievements button'),
    ('Modify:', 'GameHUD.cs - Add achievement toast notification'),
]
for label, desc in files:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(label).bold = True
    p.add_run(f' {desc}')

doc.add_page_break()

# ============================================================
# 3. MUTATORS
# ============================================================
doc.add_heading('3. System 2: Unlockable Mutators', level=1)

doc.add_heading('Overview', level=2)
doc.add_paragraph(
    'Game modifiers that change the rules of a run. Selected on the main menu before starting. '
    'Initially locked; unlocked by earning gold-tier achievements. Each mutator has a score '
    'multiplier (affects composite score). Multiple mutators can stack.'
)

doc.add_heading('Mutator Definitions (10 total)', level=2)
add_table(
    ['Mutator', 'Unlock Condition', 'Effect', 'Score Mult'],
    [
        ['Blood Tide', 'Gold "Orc Slayer"', 'Enemy count per day +50%, loot drop +30%', '1.3x'],
        ['Glass Fortress', 'Gold "Stand Your Ground"', 'Wall HP halved, defenders deal +50% damage', '1.4x'],
        ['Lone Ballista', 'Gold "Sharpshooter"', 'Cannot hire defenders; ballista dmg +100%, rate +50%', '1.8x'],
        ['Golden Horde', 'Gold "Treasure Hoarder"', 'All costs doubled, loot value tripled', '1.2x'],
        ['Night Terrors', 'Gold "Nightwatch"', 'Full enemy waves also spawn at night (no safe period)', '1.6x'],
        ['Skeleton Crew', 'Gold "Specialist"', 'Start with 1 menial; refugees arrive 2x faster', '1.3x'],
        ['Iron March', 'Gold "Against All Odds"', 'Enemies have +30% speed, cannot be slowed', '1.5x'],
        ['Bounty Hunter', 'Gold "Boss Hunter"', 'Boss spawns every 5 days; boss loot = 50g', '1.4x'],
        ['Pacifist Run', 'Gold "Army Builder"', 'Defenders cannot attack; only repair and body-block', '2.0x'],
        ['Chaos Modifiers', 'Gold "Score Legend"', 'Daily event multipliers are 2x stronger', '1.3x'],
    ]
)

doc.add_heading('Mutator Stacking Rules', level=2)
stacking = [
    'Score multipliers multiply together: Blood Tide (1.3x) + Glass Fortress (1.4x) = 1.82x total',
    'Cap total score multiplier at 5.0x to prevent absurd scores',
    'Incompatibility: Lone Ballista and Pacifist Run cannot be combined (contradictory)',
    'All other combinations are valid',
]
for s in stacking:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('Data Model', level=2)
doc.add_paragraph(
    'Each mutator stores: id (string), name, description, required achievement ID, '
    'score multiplier (float), and an array of incompatible mutator IDs.'
)

doc.add_heading('Persistence', level=2)
items = [
    'Unlocked mutators stored in PlayerPrefs key "UnlockedMutators" (JSON string array)',
    'Active mutators for current run stored in GameSettings static class (cleared on run end)',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('UI Integration', level=2)
ui_items = [
    ('Main Menu:', 'New "Mutators" button opens a toggle-list panel. Each mutator shows name, description, score multiplier, lock/unlock status. Locked mutators show the achievement needed. Toggle on/off; combined multiplier at bottom. Incompatible mutators auto-disable.'),
    ('Game HUD:', 'Active mutators shown as small text below the timer during gameplay'),
    ('Game Over:', 'Score breakdown shows base score and mutator multiplier applied'),
]
for label, desc in ui_items:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(label).bold = True
    p.add_run(f' {desc}')

doc.add_heading('Files to Create / Modify', level=2)
files = [
    ('New:', 'MutatorManager.cs - Singleton, holds active mutator list, provides multiplier queries'),
    ('New:', 'MutatorUI.cs - Main menu mutator selection panel'),
    ('Modify:', 'GameSettings.cs - Store active mutators for current run'),
    ('Modify:', 'EnemySpawnManager.cs - Apply mutator effects (spawn count, speed, night spawning)'),
    ('Modify:', 'Enemy.cs - Apply HP/speed mutator modifiers at spawn'),
    ('Modify:', 'Defender.cs - Apply damage mutator modifiers'),
    ('Modify:', 'DailyEventManager.cs - Apply Chaos Modifiers mutator (2x event strength)'),
    ('Modify:', 'GameOverScreen.cs - Show mutator score multiplier in score breakdown'),
    ('Modify:', 'MainMenuManager.cs - Add Mutators button'),
    ('Modify:', 'Wall.cs - Apply Glass Fortress HP modifier'),
]
for label, desc in files:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(label).bold = True
    p.add_run(f' {desc}')

doc.add_page_break()

# ============================================================
# 4. META-PROGRESSION
# ============================================================
doc.add_heading('4. System 3: Meta-Progression (Legacy Ranks)', level=1)

doc.add_heading('Overview', level=2)
doc.add_paragraph(
    'A prestige system where cumulative play earns "Legacy Points" that translate into '
    'small permanent bonuses. Designed to reward continued play without trivializing the game. '
    'Bonuses are modest (5-15% range at max) and take many runs to reach.'
)

doc.add_heading('Legacy Points Formula', level=2)
p = doc.add_paragraph()
p.add_run('Legacy Points = floor(CompositeScore / 1000)').bold = True
doc.add_paragraph('Example: A score of 8,500 earns 8 Legacy Points.')

doc.add_heading('Legacy Rank Tiers', level=2)
add_table(
    ['Rank', 'Points Required', 'Cumulative Bonus'],
    [
        ['Rank 0 (Recruit)', '0', 'None'],
        ['Rank 1 (Militia)', '10', '+5 starting gold'],
        ['Rank 2 (Veteran)', '30', '+5% menial movement speed'],
        ['Rank 3 (Captain)', '60', '+1 starting menial'],
        ['Rank 4 (Commander)', '100', '+5% ballista damage'],
        ['Rank 5 (Marshal)', '150', '+10 starting gold (total +15)'],
        ['Rank 6 (Warden)', '220', '+5% defender attack speed'],
        ['Rank 7 (Champion)', '300', '+5% wall max HP'],
        ['Rank 8 (Overlord)', '400', '+1 starting menial (total +2)'],
        ['Rank 9 (Legend)', '550', '+10% loot value'],
        ['Rank 10 (Mythic)', '750', 'All bonuses doubled'],
    ]
)

doc.add_heading('Design Constraints', level=2)
constraints = [
    'Bonuses are additive (not multiplicative with each other)',
    'Rank 10 doubles all bonuses but still caps at modest levels',
    'At max rank: +30g starting gold, +10% menial speed, +2 starting menials, +10% ballista damage, +10% defender attack speed, +10% wall HP, +20% loot value',
    'These bonuses make the game slightly easier but do not break balance \u2014 a Nightmare run is still very challenging at max rank',
    'Legacy points are never lost \u2014 they only accumulate',
]
for c in constraints:
    doc.add_paragraph(c, style='List Bullet')

doc.add_heading('Persistence', level=2)
doc.add_paragraph(
    'PlayerPrefs key "LegacyPoints" (int) \u2014 cumulative total. '
    'Current rank is derived from points at runtime (no need to store separately).'
)

doc.add_heading('UI Integration', level=2)
ui_items = [
    ('Main Menu:', 'Legacy Rank display near difficulty area. Shows current rank title, progress bar to next rank, tooltip listing all active bonuses.'),
    ('Game Over Screen:', 'Show "+X Legacy Points earned" with running total and progress to next rank.'),
    ('In-Game:', 'No in-game display needed (bonuses are silent/passive).'),
]
for label, desc in ui_items:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(label).bold = True
    p.add_run(f' {desc}')

doc.add_heading('Files to Create / Modify', level=2)
files = [
    ('New:', 'LegacyProgressionManager.cs - Singleton, calculates rank from points, provides bonus queries'),
    ('Modify:', 'GameSettings.cs - Apply legacy bonuses to starting gold, starting menials'),
    ('Modify:', 'GameManager.cs - Apply starting resource bonuses from legacy rank'),
    ('Modify:', 'Menial.cs / MenialManager.cs - Apply menial speed bonus'),
    ('Modify:', 'Ballista.cs - Apply ballista damage bonus'),
    ('Modify:', 'Defender.cs - Apply attack speed bonus'),
    ('Modify:', 'Wall.cs - Apply wall HP bonus'),
    ('Modify:', 'TreasurePickup.cs - Apply loot value bonus'),
    ('Modify:', 'GameOverScreen.cs - Show legacy points earned'),
    ('Modify:', 'MainMenuManager.cs - Show legacy rank display'),
]
for label, desc in files:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(label).bold = True
    p.add_run(f' {desc}')

doc.add_page_break()

# ============================================================
# 5. STATISTICS DASHBOARD
# ============================================================
doc.add_heading('5. System 4: Statistics Dashboard', level=1)

doc.add_heading('Overview', level=2)
doc.add_paragraph(
    'A comprehensive stats screen accessible from the main menu showing lifetime and per-run '
    'statistics. Helps players track improvement and discover playstyle patterns.'
)

doc.add_heading('Lifetime Totals (cumulative across all runs)', level=2)
stats = [
    'Total runs played',
    'Total days survived',
    'Total enemies killed (broken down by type: Grunts, Bow Orcs, Trolls, Goblins, Cannoneers, Bosses)',
    'Total gold earned',
    'Total gold spent',
    'Total defenders hired (by type: Engineer, Pikeman, Crossbowman, Wizard)',
    'Total menials lost',
    'Total walls built',
    'Total wall HP repaired',
    'Total vegetation cleared',
    'Total refugees saved',
    'Total ballista shots fired',
]
for s in stats:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('Personal Records (best single-run values)', level=2)
records = [
    'Longest run (days survived)',
    'Highest composite score',
    'Most kills in one run',
    'Most gold earned in one run',
    'Most defenders alive simultaneously',
    'Fastest boss kill (game time when first boss killed)',
    'Longest run per difficulty (Easy / Normal / Hard / Nightmare)',
]
for r in records:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading('Computed Stats (derived from raw data)', level=2)
computed = [
    'Average days survived per run',
    'Average kills per run',
    'Average gold per run',
    'Kill/death ratio (kills per menials lost)',
    'Favorite defender type (most hired lifetime)',
    'Most dangerous enemy type (type that appears in most game-over runs)',
    'Average score trend (last 5 runs vs previous 5 runs \u2014 improving?)',
]
for c in computed:
    doc.add_paragraph(c, style='List Bullet')

doc.add_heading('Data Model', level=2)
doc.add_paragraph(
    'A LifetimeStats class stores all totals (ints), kill/defender arrays indexed by enum, '
    'and personal records. Serialized as JSON to PlayerPrefs key "LifetimeStats". '
    'Updated at end of each run by merging run stats into lifetime stats.'
)

doc.add_heading('UI Layout', level=2)
doc.add_paragraph(
    '"Statistics" button on Main Menu opens a fullscreen panel with 5 tabs:'
)
tabs = [
    ('Overview Tab:', 'Key lifetime stats in a grid (total runs, total kills, total gold, total days)'),
    ('Records Tab:', 'Personal best records with difficulty and date achieved'),
    ('Combat Tab:', 'Kill breakdown by enemy type, ballista shot stats'),
    ('Economy Tab:', 'Gold earned/spent ratios, menial efficiency stats'),
    ('Trends Tab:', 'Last 10 run scores displayed as a simple bar chart (using UI Images for bars)'),
]
for label, desc in tabs:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(label).bold = True
    p.add_run(f' {desc}')

doc.add_heading('Files to Create / Modify', level=2)
files = [
    ('New:', 'LifetimeStatsManager.cs - Singleton, merges run stats, provides lifetime queries'),
    ('New:', 'StatsUI.cs - Main menu statistics panel with tabbed layout'),
    ('Modify:', 'RunStatsTracker.cs - Track additional stats (ballista shots, wall repairs, vegetation cleared, refugees, kills by type, gold spent)'),
    ('Modify:', 'GameOverScreen.cs - Merge run stats into lifetime stats'),
    ('Modify:', 'MainMenuManager.cs - Add Statistics button'),
    ('Modify:', 'Ballista.cs - Increment shot counter on fire'),
    ('Modify:', 'Wall.cs - Track HP repaired'),
    ('Modify:', 'Enemy.cs - Report enemy type on death for type-specific kill tracking'),
]
for label, desc in files:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(label).bold = True
    p.add_run(f' {desc}')

doc.add_page_break()

# ============================================================
# 6. SYSTEM INTERCONNECTIONS
# ============================================================
doc.add_heading('6. System Interconnections', level=1)

doc.add_paragraph('The four systems feed into each other through a common data pipeline:')

doc.add_paragraph()
p = doc.add_paragraph()
p.style = doc.styles['Normal']
p.paragraph_format.left_indent = Inches(0.5)
p.add_run(
    'Gameplay Events\n'
    '    \u2502\n'
    '    v\n'
    'RunStatsTracker (per-run data)\n'
    '    \u2502\n'
    '    \u251c\u2500\u2500\u2500> AchievementManager (check thresholds, unlock tiers)\n'
    '    \u2502         \u2502\n'
    '    \u2502         \u2514\u2500\u2500\u2500> MutatorManager (gold achievements unlock mutators)\n'
    '    \u2502\n'
    '    \u251c\u2500\u2500\u2500> LifetimeStatsManager (merge into lifetime totals)\n'
    '    \u2502\n'
    '    \u251c\u2500\u2500\u2500> LegacyProgressionManager (earn legacy points from score)\n'
    '    \u2502\n'
    '    \u2514\u2500\u2500\u2500> RunHistoryManager (existing, save run record)\n'
).font.name = 'Consolas'

doc.add_heading('Flow at Game Over', level=2)
flow = [
    'RunStatsTracker finalizes run stats',
    'AchievementManager evaluates all achievement conditions, fires unlock events',
    'LifetimeStatsManager merges run stats into lifetime data',
    'LegacyProgressionManager adds legacy points based on score',
    'RunHistoryManager saves run to top-20 leaderboard',
    'GameOverScreen displays all of the above',
]
for i, step in enumerate(flow, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_page_break()

# ============================================================
# 7. IMPLEMENTATION ORDER
# ============================================================
doc.add_heading('7. Implementation Order (Recommended)', level=1)

doc.add_paragraph(
    'The systems should be implemented in the following order to minimize '
    'dependencies and maximize incremental value:'
)

order = [
    ('1. Statistics Dashboard',
     'Requires extending RunStatsTracker with additional tracked stats (kills by type, '
     'ballista shots, wall repairs, etc.). This enriched stat tracking is foundational '
     'for the achievement system.'),
    ('2. Achievements',
     'Depends on the enriched stats tracking from Step 1. Provides the unlock mechanism '
     'for mutators in Step 4.'),
    ('3. Meta-Progression (Legacy Ranks)',
     'A relatively simple standalone system. Only depends on the composite scoring that '
     'already exists. Can be built in parallel with achievements.'),
    ('4. Mutators',
     'Depends on achievements for unlock conditions. Has the most gameplay modifications '
     '(touching enemy spawning, defender behavior, wall HP, etc.) so should be implemented '
     'last when all other systems are stable.'),
]
for title, desc in order:
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    doc.add_paragraph(desc)

# ============================================================
# SAVE
# ============================================================
output_path = '/home/user/Orc/Replayability_Design_Document.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
